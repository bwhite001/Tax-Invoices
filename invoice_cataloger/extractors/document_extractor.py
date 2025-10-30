"""
Document Text Extraction for Word and Excel files
"""
from pathlib import Path
from typing import Optional, Tuple

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False


class DocumentExtractor:
    """Extract text from Word and Excel documents"""
    
    def extract_from_word(self, doc_path: Path) -> Tuple[Optional[str], str]:
        """
        Extract text from Word document (.doc, .docx)
        
        Returns:
            Tuple[Optional[str], str]: (extracted_text, method_used)
        """
        doc_path = Path(doc_path)
        
        if not doc_path.exists():
            return None, "File not found"
        
        # For .docx files, try python-docx first
        if doc_path.suffix.lower() == '.docx' and DOCX_AVAILABLE:
            text, success = self._extract_docx_with_python_docx(doc_path)
            if success:
                return text, "python-docx"
        
        # Fallback to COM (Windows only, works for both .doc and .docx)
        if WIN32COM_AVAILABLE:
            text, success = self._extract_word_with_com(doc_path)
            if success:
                return text, "Word COM"
        
        return None, "No Word extraction method available"
    
    def extract_from_excel(self, excel_path: Path) -> Tuple[Optional[str], str]:
        """
        Extract text from Excel document (.xls, .xlsx)
        
        Returns:
            Tuple[Optional[str], str]: (extracted_text, method_used)
        """
        excel_path = Path(excel_path)
        
        if not excel_path.exists():
            return None, "File not found"
        
        # For .xlsx files, try openpyxl first
        if excel_path.suffix.lower() == '.xlsx' and OPENPYXL_AVAILABLE:
            text, success = self._extract_xlsx_with_openpyxl(excel_path)
            if success:
                return text, "openpyxl"
        
        # Fallback to COM (Windows only, works for both .xls and .xlsx)
        if WIN32COM_AVAILABLE:
            text, success = self._extract_excel_with_com(excel_path)
            if success:
                return text, "Excel COM"
        
        return None, "No Excel extraction method available"
    
    def _extract_docx_with_python_docx(self, doc_path: Path) -> Tuple[Optional[str], bool]:
        """Extract text from .docx using python-docx"""
        try:
            doc = Document(doc_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            if text and len(text.strip()) > 10:
                return text, True
            
            return None, False
        except Exception:
            return None, False
    
    def _extract_word_with_com(self, doc_path: Path) -> Tuple[Optional[str], bool]:
        """Extract text from Word document using COM"""
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            
            doc = word.Documents.Open(str(doc_path.absolute()), ReadOnly=True)
            text = doc.Content.Text
            
            doc.Close(False)
            word.Quit()
            
            if text and len(text.strip()) > 10:
                return text, True
            
            return None, False
        except Exception:
            return None, False
    
    def _extract_xlsx_with_openpyxl(self, excel_path: Path) -> Tuple[Optional[str], bool]:
        """Extract text from .xlsx using openpyxl"""
        try:
            workbook = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
            text = ""
            
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            
            workbook.close()
            
            if text and len(text.strip()) > 10:
                return text, True
            
            return None, False
        except Exception:
            return None, False
    
    def _extract_excel_with_com(self, excel_path: Path) -> Tuple[Optional[str], bool]:
        """Extract text from Excel document using COM"""
        try:
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            
            workbook = excel.Workbooks.Open(str(excel_path.absolute()), ReadOnly=True)
            text = ""
            
            for worksheet in workbook.Worksheets:
                used_range = worksheet.UsedRange
                if used_range:
                    for row in used_range.Rows:
                        for cell in row.Cells:
                            if cell.Text:
                                text += str(cell.Text) + " "
                        text += "\n"
            
            workbook.Close(False)
            excel.Quit()
            
            if text and len(text.strip()) > 10:
                return text, True
            
            return None, False
        except Exception:
            return None, False
    
    @staticmethod
    def get_available_methods() -> dict:
        """Get information about available extraction methods"""
        return {
            'python-docx': DOCX_AVAILABLE,
            'openpyxl': OPENPYXL_AVAILABLE,
            'Win32COM': WIN32COM_AVAILABLE,
        }
